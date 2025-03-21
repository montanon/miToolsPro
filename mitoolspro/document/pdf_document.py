import re
from itertools import accumulate
from math import isclose
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from pandas import DataFrame
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTAnno, LTChar, LTTextBoxHorizontal, LTTextLineHorizontal
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from mitoolspro.exceptions import ArgumentValueError
from mitoolspro.files.document.document_structure import (
    BBox,
    Box,
    Char,
    Document,
    Line,
    Page,
    Run,
)

FILEPATH = Path(__file__).parent

FONTS = FILEPATH.parent / "fonts"

if FONTS.exists():
    NORMAL_FONT = FONTS / "arialmt.ttf"
    BOLD_FONT = FONTS / "Arial Bold.ttf"
    ITALIC_FONT = FONTS / "Arial Italic.ttf"
    BOLD_ITALIC_FONT = FONTS / "Arial Bold Italic.ttf"
    pdfmetrics.registerFont(TTFont("Arial", NORMAL_FONT))
    pdfmetrics.registerFont(TTFont("Arial-Bold", BOLD_FONT))
    pdfmetrics.registerFont(TTFont("Arial-Italic", ITALIC_FONT))
    pdfmetrics.registerFont(TTFont("Arial-BoldItalic", BOLD_ITALIC_FONT))


FONT_MAPPING = {
    "normal": "Arial",
    "bold": "Arial-Bold",
    "italic": "Arial-Italic",
    "bold-italic": "Arial-BoldItalic",
}


def _select_font(fontname):
    fname = fontname.lower()
    if "bold" in fname and ("italic" in fname or "oblique" in fname):
        return FONT_MAPPING["bold-italic"]
    elif "bold" in fname:
        return FONT_MAPPING["bold"]
    elif "italic" in fname or "oblique" in fname:
        return FONT_MAPPING["italic"]
    else:
        return FONT_MAPPING["normal"]


def get_char_properties(char: str, fontname: str, size: float, chars_data: DataFrame):
    properties = chars_data[
        (chars_data["fontname"] == fontname) & (chars_data["size"] == size)
    ]
    return properties[properties["text"] == char]


def center_positions(positions: List[float], center: float):
    current_center = (positions[0] + positions[-1]) / 2
    shift = center - current_center
    return [position + shift for position in positions]


def create_run_in_bbox(
    text: str, fontname: str, size: float, bbox: BBox, chars_data: DataFrame
):
    xcenter = bbox.xcenter
    widths = [
        get_char_properties(char, fontname, size, chars_data)["width"].item()
        for char in text
    ]
    xpositions = list(accumulate([width for width in widths]))
    xpositions = center_positions(xpositions, xcenter)
    chars = [
        Char(
            text=char,
            fontname=fontname,
            size=size,
            x0=xposition,
            y0=bbox.y0,
            x1=xposition + width,
            y1=bbox.y1,
        )
        for char, xposition, width in zip(text, xpositions, widths)
    ]
    return Run.from_chars(chars, fontname=fontname, size=size)


def _create_doc(
    top_margin: float = 2.47,
    bottom_margin: float = 1.31,
    left_margin: float = 1.5,
    right_margin: float = 1.5,
    header_distance: float = 0,
    footer_distance: float = 0.96,
):
    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Cm(top_margin)
    section.bottom_margin = Cm(bottom_margin)
    section.left_margin = Cm(left_margin)
    section.right_margin = Cm(right_margin)
    section.header_distance = Cm(header_distance)
    section.footer_distance = Cm(footer_distance)
    return doc


def _preprocess_pdf_text(text: str):
    return re.sub(r"\s+", " ", text.replace("\n", " ")).strip()


def merge_runs(runs: List):
    if not runs:
        return []
    merged_runs = []
    current_run = runs[0]
    for next_run in runs[1:]:
        same_font = next_run.fontname == current_run.fontname
        same_size = isclose(next_run.size, current_run.size, abs_tol=0.01)
        if same_font and same_size:
            current_run = current_run + next_run
        else:
            merged_runs.append(current_run)
            current_run = next_run
    merged_runs.append(current_run)
    return merged_runs


def sections_to_doc(
    doc: DocxDocument,
    sections: List[List[Run]],
    has_title: bool = False,
):
    if has_title:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(sections[0][0].text.replace("\n", " "))
        title_run.font.size = Pt(sections[0][0].size)
        title_run.font.name = "Arial MT"
        if sections[0][0].is_bold():
            title_run.bold = True
        if sections[0][0].is_italic():
            title_run.italic = True
        title_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sections = sections[1:]

    n_jump = 2
    items_indeces = find_item_runs(sections)
    items_indeces = [items_indeces[0] - n_jump] + items_indeces
    last_item_indeces = [
        items_indeces[i]
        for i in range(len(items_indeces) - 1)
        if items_indeces[i + 1] - items_indeces[i] > n_jump
    ]
    for idx in last_item_indeces:
        last_item_idx = items_indeces.index(idx)
        items_indeces = (
            items_indeces[: last_item_idx + 1]
            + [items_indeces[last_item_idx + 1] - n_jump]
            + items_indeces[last_item_idx + 1 :]
        )
    last_item_indeces.append(items_indeces[-1])

    current_run = 0
    for n_section, section in enumerate(sections):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format = paragraph.paragraph_format
        for i, run in enumerate(section):
            text_run = _preprocess_pdf_text(run.text)
            docx_run = paragraph.add_run(text_run)
            docx_run.font.size = Pt(run.size)
            docx_run.font.name = "Arial MT"
            if run.is_bold():
                docx_run.bold = True
            if run.is_italic():
                docx_run.italic = True
            docx_run.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if current_run in items_indeces:
                if current_run not in last_item_indeces:
                    p_format.space_after = Pt(0)
            current_run += 1
    return doc


def runs_to_sections(runs: List[Run]):
    sections = []
    current_section = []
    for run in runs:
        if run.is_bold():
            if current_section:
                sections.append(current_section)
            current_section = [run]
        else:
            current_section.append(run)
    if current_section:
        sections.append(current_section)
    return sections


def iterate_all_runs(sections):
    run_index = 0
    for elem in sections:
        if isinstance(elem, list):
            for run in elem:
                yield run, run_index
                run_index += 1
        else:
            yield elem, run_index
            run_index += 1


def find_item_runs(sections):
    item_indices = []
    item_pattern = re.compile(r"^(?:Item\s+\d+:|[a-z][\.\)]|\d+[\.\)])", re.IGNORECASE)
    for run, idx in iterate_all_runs(sections):
        if run.is_bold() and item_pattern.match(run.text):
            item_indices.append(idx)
    return item_indices


def insert_values_after_jumps(lst, n):
    if not lst or len(lst) < 2:
        return lst
    result = []
    for i in range(len(lst) - 1):
        result.append(lst[i])
        if lst[i + 1] - lst[i] > n:
            result.append(lst[i] + n)
    result.append(lst[-1])
    return result


def rewrite_pdf_char_by_char(doc, output_path):
    if not doc.pages:
        raise ArgumentValueError("Document has no pages")
    first_page = doc.pages[0]
    page_width = first_page.width
    page_height = first_page.height
    c = canvas.Canvas(str(output_path), pagesize=(page_width, page_height))
    for page in doc.pages:
        c.setPageSize((page.width, page.height))

        for box in page.boxes:
            for line in box.lines:
                for run in line.runs:
                    font_used = _select_font(run.fontname)
                    for char in run.chars:
                        if char.x0 != 0 or char.y0 != 0:
                            c.setFont(font_used, char.size)
                            c.setFont(font_used, char.size)
                            c.drawString(char.x0, char.y0, char.text)
        c.showPage()
    c.save()


def center_runs_vertically(runs: List[Run], reference_y: float, step: int = 4):
    sizes = [run.size for run in runs]
    ypositions = []
    current_pos = 0
    current_step = step
    for size in sizes:
        ypositions.append(current_pos)
        current_pos += current_step + size
        current_step += step
    ypositions = ypositions[::-1]
    ypositions = center_positions(ypositions, reference_y)
    for i, run in enumerate(runs):
        run.chars = [
            Char(
                text=char.text,
                fontname=char.fontname,
                size=char.size,
                x0=char.x0,
                y0=ypositions[i],
                x1=char.x1,
                y1=ypositions[i] + char.size,
            )
            for char in run.chars
        ]
    return runs

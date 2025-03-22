import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from reportlab.pdfgen import canvas

from mitoolspro.document.document_structure import (
    Document,
    Image,
    Line,
)
from mitoolspro.document.fonts import select_font
from mitoolspro.exceptions import ArgumentValueError


def write_pdf(doc: Document, output_path: Path):
    if not doc.pages:
        raise ArgumentValueError("Document has no pages")

    first_page = doc.pages[0]
    c = canvas.Canvas(str(output_path), pagesize=(first_page.width, first_page.height))

    for page in doc.pages:
        c.setPageSize((page.width, page.height))

        for box in page.boxes:
            for line in box.get_all_lines():
                for run in line.runs:
                    font_used = select_font(run.fontname)
                    for char in run.chars:
                        if char.bbox.x0 != 0 or char.bbox.y0 != 0:
                            c.setFont(font_used, char.size)
                            c.drawString(char.bbox.x0, char.bbox.y0, char.text)

        c.showPage()
    c.save()


def _create_docx(
    top_margin: float = 2.47,
    bottom_margin: float = 1.31,
    left_margin: float = 1.5,
    right_margin: float = 1.5,
    header_distance: float = 0,
    footer_distance: float = 0.96,
) -> DocxDocument:
    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Cm(top_margin)
    section.bottom_margin = Cm(bottom_margin)
    section.left_margin = Cm(left_margin)
    section.right_margin = Cm(right_margin)
    section.header_distance = Cm(header_distance)
    section.footer_distance = Cm(footer_distance)
    return doc


def write_docx(doc: Document, output_path: Path):
    docx = _create_docx()

    for page in doc.pages:
        for box in page.boxes:
            for element in box.elements:
                if isinstance(element, Line):
                    paragraph = docx.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in element.runs:
                        text_run = re.sub(
                            r"\s+", " ", run.text.replace("\n", " ")
                        ).strip()
                        docx_run = paragraph.add_run(text_run)
                        docx_run.font.size = Pt(run.size)
                        docx_run.font.name = "Arial MT"
                        docx_run.bold = run.is_bold()
                        docx_run.italic = run.is_italic()
                elif isinstance(element, Image):
                    if not element.stream:
                        continue
                    paragraph = docx.add_paragraph()
                    run = paragraph.add_run()
                    image_stream = BytesIO(element.stream)
                    width_inches = element.bbox.width / 72  # Convert from pt to inch
                    run.add_picture(image_stream, width=Inches(width_inches))

    docx.save(output_path)
    return docx

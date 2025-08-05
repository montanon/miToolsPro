import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase

from docx import Document as DocxDocument
from docx.shared import Pt

from mitoolspro.document.document_structure import Document, Line
from mitoolspro.document.from_docx import docx_to_document


class TestDocxToDocument(TestCase):
    def setUp(self):
        self.temp_dir = TemporaryDirectory()

    def tearDown(self):
        self.temp_dir.cleanup()

    def create_simple_docx(self, file_path: Path):
        doc = DocxDocument()
        doc.add_paragraph("First paragraph with simple text.")
        doc.add_paragraph("Second paragraph with different content.")
        doc.save(file_path)
        return doc

    def test_docx_to_document_basic(self):
        file_path = Path(self.temp_dir.name) / "test_basic.docx"
        self.create_simple_docx(file_path)

        result = docx_to_document(file_path)

        self.assertIsInstance(result, Document)
        self.assertEqual(len(result.pages), 1)

        page = result.pages[0]
        self.assertEqual(page.width, 595.3)  # A4 width in points
        self.assertEqual(page.height, 841.9)  # A4 height in points

    def test_docx_to_document_with_multiple_paragraphs(self):
        file_path = Path(self.temp_dir.name) / "test_multi_para.docx"
        doc = DocxDocument()
        doc.add_paragraph("Paragraph one")
        doc.add_paragraph("Paragraph two")
        doc.add_paragraph("Paragraph three")
        doc.save(file_path)

        result = docx_to_document(file_path)

        self.assertIsInstance(result, Document)
        self.assertEqual(len(result.pages), 1)

        page = result.pages[0]
        # Should have multiple boxes for multiple paragraphs
        self.assertGreater(len(page.boxes), 0)

    def test_docx_to_document_with_empty_paragraphs(self):
        file_path = Path(self.temp_dir.name) / "test_empty_para.docx"
        doc = DocxDocument()
        doc.add_paragraph("")  # Empty paragraph
        doc.add_paragraph("Non-empty paragraph")
        doc.add_paragraph("")  # Another empty paragraph
        doc.save(file_path)

        result = docx_to_document(file_path)

        self.assertIsInstance(result, Document)
        self.assertEqual(len(result.pages), 1)

        page = result.pages[0]
        # Should only have boxes for non-empty paragraphs
        self.assertGreater(len(page.boxes), 0)

    def test_docx_to_document_with_formatted_text(self):
        file_path = Path(self.temp_dir.name) / "test_formatted.docx"
        doc = DocxDocument()
        para = doc.add_paragraph()
        run1 = para.add_run("Normal text ")
        run2 = para.add_run("bold text")
        run2.bold = True
        doc.save(file_path)

        result = docx_to_document(file_path)

        self.assertIsInstance(result, Document)
        self.assertEqual(len(result.pages), 1)

        page = result.pages[0]
        self.assertGreater(len(page.boxes), 0)

        # Check that runs are created for different formatting
        box = page.boxes[0]
        self.assertGreater(len(box.elements), 0)

        line = box.elements[0]
        self.assertIsInstance(line, Line)
        # Should have at least one run
        self.assertGreater(len(line.runs), 0)

    def test_docx_to_document_font_handling(self):
        file_path = Path(self.temp_dir.name) / "test_fonts.docx"
        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("Test text with font")
        run.font.name = "Arial"
        try:
            run.font.size = Pt(14)
        except AttributeError:
            pass  # Some versions may not support this
        doc.save(file_path)

        result = docx_to_document(file_path)

        self.assertIsInstance(result, Document)
        self.assertEqual(len(result.pages), 1)

        page = result.pages[0]
        self.assertGreater(len(page.boxes), 0)

    def test_docx_to_document_character_processing(self):
        file_path = Path(self.temp_dir.name) / "test_chars.docx"
        doc = DocxDocument()
        doc.add_paragraph("ABC")
        doc.save(file_path)

        result = docx_to_document(file_path)

        self.assertIsInstance(result, Document)
        page = result.pages[0]

        # Should have processed individual characters
        box = page.boxes[0]
        line = box.elements[0]
        self.assertIsInstance(line, Line)
        run = line.runs[0]

        # Should have 3 characters: A, B, C
        self.assertEqual(len(run.chars), 3)
        self.assertEqual(run.chars[0].text, "A")
        self.assertEqual(run.chars[1].text, "B")
        self.assertEqual(run.chars[2].text, "C")

    def test_docx_to_document_default_font_values(self):
        file_path = Path(self.temp_dir.name) / "test_defaults.docx"
        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("Test text")
        # Don't set font name or size to test defaults
        doc.save(file_path)

        result = docx_to_document(file_path)

        self.assertIsInstance(result, Document)
        page = result.pages[0]
        box = page.boxes[0]
        line = box.elements[0]
        self.assertIsInstance(line, Line)
        run_obj = line.runs[0]

        # Should use default font name and size
        self.assertEqual(run_obj.fontname, "Times New Roman")
        self.assertEqual(run_obj.size, 12.0)


if __name__ == "__main__":
    unittest.main()

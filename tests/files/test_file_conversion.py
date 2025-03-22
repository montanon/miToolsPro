import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase

import PyPDF2
from docx import Document

from mitoolspro.exceptions import ArgumentValueError
from mitoolspro.files.file_conversion import (
    convert_directory_docxs_to_pdfs,
    convert_directory_files,
    convert_docx_to_pdf,
    convert_file,
)


class BaseDocTest(TestCase):
    def setUp(self):
        self.temp_dir = TemporaryDirectory()
        self.test_dir = Path(self.temp_dir.name)

        doc = Document()
        doc.add_heading("Test Heading 1", 1)
        doc.add_paragraph("Normal paragraph")
        p = doc.add_paragraph()
        p.add_run("Bold text").bold = True
        p.add_run(" and ")
        p.add_run("italic text").italic = True

        self.test_docx = self.test_dir / "test.docx"
        doc.save(self.test_docx)

    def tearDown(self):
        self.temp_dir.cleanup()


class TestConvertDocxToPdf(BaseDocTest):
    def test_basic_conversion(self):
        output_pdf = self.test_dir / "output.pdf"
        convert_docx_to_pdf(self.test_docx, output_pdf)
        self.assertTrue(output_pdf.exists())

        with open(output_pdf, "rb") as f:
            PyPDF2.PdfReader(f)

    def test_invalid_input_file(self):
        output_pdf = self.test_dir / "output.pdf"
        with self.assertRaises(Exception):
            convert_docx_to_pdf(self.test_dir / "nonexistent.docx", output_pdf)


class TestConvertFile(BaseDocTest):
    def test_successful_conversion(self):
        output_file = self.test_dir / "output.pdf"
        convert_file(self.test_docx, output_file, "pdf")
        self.assertTrue(output_file.exists())

    def test_exists_with_overwrite(self):
        output_file = self.test_dir / "exists.pdf"
        output_file.touch()

        convert_file(self.test_docx, output_file, "pdf", exist_ok=True, overwrite=True)
        self.assertTrue(output_file.exists())

    def test_exists_no_overwrite(self):
        output_file = self.test_dir / "exists.pdf"
        output_file.touch()

        convert_file(self.test_docx, output_file, "pdf", exist_ok=True, overwrite=False)

        with self.assertRaises(ArgumentValueError):
            convert_file(
                self.test_docx, output_file, "pdf", exist_ok=False, overwrite=False
            )


class TestConvertDirectoryFiles(BaseDocTest):
    def test_basic_directory_conversion(self):
        output_dir = self.test_dir / "output"
        output_dir.mkdir()

        convert_directory_files(self.test_dir, output_dir, "docx", "pdf")

        pdf_files = list(output_dir.glob("*.pdf"))
        self.assertEqual(len(pdf_files), 1)

        with open(pdf_files[0], "rb") as f:
            PyPDF2.PdfReader(f)

    def test_empty_directory(self):
        empty_dir = self.test_dir / "empty"
        empty_dir.mkdir()
        output_dir = self.test_dir / "output"
        output_dir.mkdir()

        convert_directory_files(empty_dir, output_dir, "docx", "pdf")

        self.assertEqual(len(list(output_dir.glob("*.pdf"))), 0)


class TestConvertDirectoryDocxsToPdfs(BaseDocTest):
    def test_with_output_directory(self):
        output_dir = self.test_dir / "pdf_output"
        output_dir.mkdir()

        convert_directory_docxs_to_pdfs(self.test_dir, output_dir)

        pdf_files = list(output_dir.glob("*.pdf"))
        self.assertEqual(len(pdf_files), 1)

        with open(pdf_files[0], "rb") as f:
            PyPDF2.PdfReader(f)

    def test_without_output_directory(self):
        convert_directory_docxs_to_pdfs(self.test_dir, None)

        pdf_files = list(self.test_dir.glob("*.pdf"))
        self.assertEqual(len(pdf_files), 1)

    def test_empty_directory(self):
        empty_dir = self.test_dir / "empty"
        empty_dir.mkdir()
        output_dir = self.test_dir / "output"
        output_dir.mkdir()

        convert_directory_docxs_to_pdfs(empty_dir, output_dir)
        self.assertEqual(len(list(output_dir.glob("*.pdf"))), 0)


if __name__ == "__main__":
    unittest.main()

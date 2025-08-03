import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase

from docx import Document

from mitoolspro.files.doc_handlers import read_docx_file


class TestReadDocxFile(TestCase):
    def setUp(self):
        self.temp_dir = TemporaryDirectory()

    def tearDown(self):
        self.temp_dir.cleanup()

    def create_simple_docx(self, file_path: Path):
        doc = Document()

        # Add heading
        doc.add_heading("Test Heading 1", level=1)

        # Add normal paragraph
        doc.add_paragraph("This is a normal paragraph.")

        # Add paragraph with bold and italic
        para2 = doc.add_paragraph()
        para2.add_run("This is ")
        run2 = para2.add_run("bold")
        run2.bold = True
        para2.add_run(" and ")
        run4 = para2.add_run("italic")
        run4.italic = True
        para2.add_run(" text.")

        # Add list paragraph
        doc.add_paragraph("List item", style="List Paragraph")

        # Add paragraph ending with colon
        doc.add_paragraph("Category:")

        # Add second level heading
        doc.add_heading("Test Heading 2", level=2)

        doc.save(file_path)
        return doc

    def test_read_basic_docx_file(self):
        file_path = Path(self.temp_dir.name) / "test_basic.docx"
        self.create_simple_docx(file_path)

        result = read_docx_file(file_path)

        self.assertIsInstance(result, list)
        self.assertGreater(len(result), 0)

        # Check if heading is properly formatted
        heading_found = False
        for line in result:
            if line.strip().startswith("# Test Heading 1"):
                heading_found = True
                break
        self.assertTrue(heading_found)

    def test_read_docx_with_custom_indent(self):
        file_path = Path(self.temp_dir.name) / "test_indent.docx"
        self.create_simple_docx(file_path)

        result = read_docx_file(file_path, indent="*")

        self.assertIsInstance(result, list)
        self.assertGreater(len(result), 0)

    def test_read_docx_with_path_string(self):
        file_path = Path(self.temp_dir.name) / "test_string.docx"
        self.create_simple_docx(file_path)

        result = read_docx_file(str(file_path))

        self.assertIsInstance(result, list)
        self.assertGreater(len(result), 0)

    def test_read_docx_with_empty_paragraphs(self):
        file_path = Path(self.temp_dir.name) / "test_empty.docx"
        doc = Document()

        # Add empty paragraph
        doc.add_paragraph("")

        # Add non-empty paragraph
        doc.add_paragraph("Non-empty paragraph")

        # Add another empty paragraph
        doc.add_paragraph("")

        doc.save(file_path)

        result = read_docx_file(file_path)

        # Should only return non-empty paragraphs
        self.assertEqual(len(result), 1)
        self.assertIn("Non-empty paragraph", result[0])

    def test_read_docx_formatting_detection(self):
        file_path = Path(self.temp_dir.name) / "test_formatting.docx"
        doc = Document()

        # Add paragraph with mixed formatting
        para = doc.add_paragraph()
        para.add_run("Normal ")
        bold_run = para.add_run("bold")
        bold_run.bold = True
        para.add_run(" and ")
        italic_run = para.add_run("italic")
        italic_run.italic = True
        para.add_run(" text.")

        doc.save(file_path)

        result = read_docx_file(file_path)

        self.assertEqual(len(result), 1)
        # Check for bold and italic markdown formatting
        self.assertIn("**bold**", result[0])
        self.assertIn("*italic*", result[0])

    def test_read_docx_headings_hierarchy(self):
        file_path = Path(self.temp_dir.name) / "test_headings.docx"
        doc = Document()

        doc.add_heading("Heading 1", level=1)
        doc.add_heading("Heading 2", level=2)
        doc.add_heading("Heading 3", level=3)

        doc.save(file_path)

        result = read_docx_file(file_path)

        self.assertEqual(len(result), 3)
        self.assertTrue(any("# Heading 1" in line for line in result))
        self.assertTrue(any("## Heading 2" in line for line in result))
        self.assertTrue(any("### Heading 3" in line for line in result))

    def test_read_docx_list_indentation(self):
        file_path = Path(self.temp_dir.name) / "test_lists.docx"
        doc = Document()

        doc.add_paragraph("Normal paragraph")
        doc.add_paragraph("List item 1", style="List Paragraph")
        doc.add_paragraph("List item 2", style="List Paragraph")
        doc.add_paragraph("Back to normal", style="Normal")

        doc.save(file_path)

        result = read_docx_file(file_path)

        self.assertEqual(len(result), 4)
        # List items should have indentation
        list_items = [line for line in result if "List item" in line]
        self.assertEqual(len(list_items), 2)
        for item in list_items:
            self.assertTrue(item.startswith("\t-"))


if __name__ == "__main__":
    unittest.main()
